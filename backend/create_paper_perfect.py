#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
양식 파일의 정확한 스타일과 폰트를 사용하여 완벽한 논문 생성
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
import sys

def set_two_column_section(section):
    """섹션을 2단 컬럼으로 설정"""
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0] if sectPr.xpath('./w:cols') else OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    # 양식 파일 기준: 28.35pt = 567 (20pt 단위)
    # 하지만 일반적인 2단 컬럼 간격은 0.5인치(708)이므로 양식과 일치시키기 위해 567 사용
    cols.set(qn('w:space'), '567')  # 컬럼 간격 (28.35pt, 양식 파일과 일치)
    if not sectPr.xpath('./w:cols'):
        sectPr.append(cols)

def extract_english_title(markdown_content):
    """영문 제목 추출"""
    lines = markdown_content.split('\n')
    for i, line in enumerate(lines):
        if 'Behavior-Based' in line or 'Digital Human Twin' in line:
            # 다음 줄이 영문 제목일 가능성
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if next_line and not next_line.startswith('#'):
                    return next_line
    return "Behavior-Based Digital Human Twin: A Machine Learning Framework for Personality Inference from Gameplay Data"

def parse_markdown_section(content):
    """마크다운 내용을 파싱하여 구조화된 데이터로 변환"""
    lines = content.split('\n')
    sections = []
    current_section = None
    current_subsection = None
    abstract_content = []
    in_abstract = False
    
    for i, line in enumerate(lines):
        original_line = line
        line = line.strip()
        
        if '## 초록' in line or '## Abstract' in line:
            in_abstract = True
            if current_section:
                sections.append(current_section)
            current_section = {
                'title': '요 약',
                'level': 1,
                'subsections': [],
                'content': []
            }
            continue
        elif in_abstract and line.startswith('##'):
            in_abstract = False
        
        if line.startswith('## ') and not in_abstract:
            if current_section:
                sections.append(current_section)
            current_section = {
                'title': line[3:].strip(),
                'level': 1,
                'subsections': [],
                'content': []
            }
            current_subsection = None
        elif line.startswith('### '):
            if current_subsection and current_section:
                current_section['subsections'].append(current_subsection)
            current_subsection = {
                'title': line[4:].strip(),
                'level': 2,
                'content': []
            }
        elif current_subsection and current_section:
            current_subsection['content'].append(original_line)
        elif current_section:
            current_section['content'].append(original_line)
    
    if current_subsection and current_section:
        current_section['subsections'].append(current_subsection)
    if current_section:
        sections.append(current_section)
    
    return sections

def format_references(markdown_content):
    """참고문헌 포맷팅"""
    refs = []
    in_refs = False
    
    for line in markdown_content.split('\n'):
        if '## 참고문헌' in line or '## 참고 문헌' in line:
            in_refs = True
            continue
        elif in_refs and line.startswith('##'):
            break
        elif in_refs and line.strip():
            # [1], [2] 형식 또는 1., 2. 형식 모두 처리
            if line.strip().startswith('[') and ']' in line:
                # [1] 형식: [1] 제거
                ref = re.sub(r'^\[\d+\]\s*', '', line.strip())
                refs.append(ref)
            elif line.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                # 1. 형식: 번호 제거
                ref = re.sub(r'^\d+\.\s*', '', line.strip())
                refs.append(ref)
    
    return refs

def add_formatted_paragraph(doc, text, style_name='바탕글'):
    """양식에 맞는 단락 추가"""
    if text.startswith('```'):
        return doc.add_paragraph()
    
    para = doc.add_paragraph(style=style_name)
    
    # **강조** 처리
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)
    
    return para

def create_perfect_paper(template_path, markdown_path, output_path):
    """양식 파일의 정확한 스타일을 사용하여 완벽한 논문 생성"""
    
    print(f"양식 파일 읽는 중: {template_path}")
    doc = Document(template_path)
    
    # 기존 내용 모두 삭제
    body = doc._body._body
    for para in list(body):
        body.remove(para)
    
    # 페이지 설정 (양식 파일과 동일하게)
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # 마크다운 파일 읽기
    print(f"마크다운 파일 읽는 중: {markdown_path}")
    with open(markdown_path, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # 메타데이터 추출
    title_match = re.search(r'^# (.+)$', markdown_content, re.MULTILINE)
    title = title_match.group(1) if title_match else "행동 기반 디지털 휴먼 트윈"
    
    english_title = extract_english_title(markdown_content)
    
    # 저자 정보 추출
    author_match = re.search(r'\*\*저자\*\*: (.+)', markdown_content)
    author_name = author_match.group(1).replace('*', '').strip() if author_match else "유준석"
    
    affiliation_match = re.search(r'\*\*소속\*\*: (.+)', markdown_content)
    affiliation = affiliation_match.group(1).replace('*', '').strip() if affiliation_match else "한신대학교 공공인재학부"
    
    email_match = re.search(r'\*\*이메일\*\*: (.+)', markdown_content)
    email = email_match.group(1).strip() if email_match else "dbwnstjr1973@hs.ac.kr"
    
    # ===== 제목 영역 (1단) =====
    # 최상단 이미지 추가 (PDF 양식 기준: 위치 y=781.57pt, 크기 24.35×23.51pt)
    # 이미지는 헤더 텍스트 위 또는 왼쪽에 배치
    base_dir = os.path.dirname(os.path.abspath(__file__))
    possible_image_paths = [
        os.path.join(base_dir, '..', 'docs', 'top_logo.png'),
        os.path.join(base_dir, '..', 'docs', 'top_logo.jpg'),
        os.path.join(base_dir, '..', 'docs', 'logo.png'),
        os.path.join(base_dir, '..', 'docs', 'logo.jpg'),
        os.path.join(base_dir, '..', 'docs', 'top_image.png'),
    ]
    
    image_added = False
    image_path = None
    for img_path in possible_image_paths:
        if os.path.exists(img_path):
            image_path = img_path
            break
    
    # 이미지가 있으면 별도 단락으로 추가 (헤더 위)
    if image_path:
        try:
            # 이미지 전용 단락 생성 (가운데 정렬)
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            # 이미지 크기: 양식 기준 24.35×23.51pt = 약 0.86cm × 0.83cm
            run.add_picture(image_path, width=Cm(0.86), height=Cm(0.83))
            image_added = True
            print(f"✓ 최상단 이미지 추가: {image_path}")
            doc.add_paragraph()  # 빈 줄
        except Exception as e:
            print(f"⚠️ 이미지 추가 실패: {e}")
            print(f"   경로: {image_path}")
    
    # 헤더 텍스트 (바탕글 스타일, 9pt)
    header_para = doc.add_paragraph("Hanshin University 2022 ABC Camp 논문집 제#권 제#호", style='바탕글')
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header_para.runs:
        run.font.size = Pt(9)
        run.font.name = '한양신명조'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '한양신명조')
    
    if not image_added:
        print("ℹ️ 최상단 이미지 파일을 찾을 수 없습니다.")
        print("   PDF에서 추출한 이미지를 docs/top_logo.png로 저장하거나,")
        print("   이미지 파일을 docs/ 폴더에 추가하세요.")
    
    doc.add_paragraph()  # 빈 줄
    
    # 한국어 제목 (국문제목 스타일: 한양신명조, 16pt, 굵게, 가운데)
    title_para = doc.add_paragraph(title, style='국문제목')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.size = Pt(16)
        run.font.name = '한양신명조'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '한양신명조')
        run.bold = True
    
    doc.add_paragraph()  # 빈 줄
    
    # 영문 제목 (영문제목 스타일: 한양신명조, 14pt, 굵게, 가운데)
    english_title_para = doc.add_paragraph(english_title, style='영문제목')
    english_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in english_title_para.runs:
        run.font.size = Pt(14)
        run.font.name = '한양신명조'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '한양신명조')
        run.bold = True
    
    doc.add_paragraph()  # 빈 줄
    
    # 저자 정보 (저자 스타일: HCI Poppy, 가운데)
    author_para = doc.add_paragraph(f"{author_name}*", style='저자')
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in author_para.runs:
        run.font.name = 'HCI Poppy'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'HCI Poppy')
    
    # 소속 정보 (저자 스타일)
    affiliation_para = doc.add_paragraph(f"*{affiliation}", style='저자')
    affiliation_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in affiliation_para.runs:
        run.font.name = 'HCI Poppy'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'HCI Poppy')
    
    # 이메일 정보 (저자 스타일, 선택적)
    if email:
        email_para = doc.add_paragraph(email, style='저자')
        email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in email_para.runs:
            run.font.name = 'HCI Poppy'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'HCI Poppy')
    
    doc.add_paragraph()  # 빈 줄
    
    # ===== 본문 영역 =====
    # 섹션 파싱
    sections = parse_markdown_section(markdown_content)
    
    # 요약 섹션은 1단으로
    abstract_section = None
    other_sections = []
    for section in sections:
        if section['title'] == '요 약':
            abstract_section = section
        else:
            other_sections.append(section)
    
    # 요약 섹션 추가 (1단, 요약 스타일: 한양신명조, 9pt)
    if abstract_section:
        doc.add_paragraph("요 약", style='요약')
        for content_line in abstract_section.get('content', []):
            if content_line.strip() and not content_line.startswith('```'):
                para = add_formatted_paragraph(doc, content_line, style_name='요약')
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = '한양신명조'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '한양신명조')
        doc.add_paragraph()  # 빈 줄
    
    # 본문 섹션들을 위한 새 섹션 생성 (2단 컬럼)
    doc.add_section()
    section = doc.sections[-1]
    
    # 페이지 설정
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    
    # 2단 컬럼 설정
    set_two_column_section(section)
    
    # 각 섹션 추가 (2단 컬럼)
    section_num = 1
    for section_data in other_sections:
        if section_data['title'] == '요 약':
            continue
        
        # 섹션 제목 (1. 스타일: 한양신명조, 굵게)
        section_title = f"{section_num}. {section_data['title']}"
        title_para = doc.add_paragraph(section_title, style='1.')
        for run in title_para.runs:
            run.font.name = '한양신명조'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '한양신명조')
            run.bold = True
        
        # 서브섹션 처리
        subsection_num = 1
        for subsection in section_data.get('subsections', []):
            subsection_title = f"{section_num}.{subsection_num} {subsection['title']}"
            sub_para = doc.add_paragraph(subsection_title, style='1.1')
            for run in sub_para.runs:
                run.font.size = Pt(9)
                run.font.name = '한양견명조'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '한양견명조')
            
            for content_line in subsection.get('content', []):
                if content_line.strip():
                    para = add_formatted_paragraph(doc, content_line, style_name='바탕글')
                    for run in para.runs:
                        run.font.size = Pt(9)
                        run.font.name = '한양신명조'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '한양신명조')
            
            subsection_num += 1
        
        # 섹션 본문 내용 (바탕글 스타일: 한양신명조, 9pt)
        for content_line in section_data.get('content', []):
            if content_line.strip() and not content_line.startswith('```'):
                content_line = re.sub(r'\[(\d+)\]', r'[\1]', content_line)
                para = add_formatted_paragraph(doc, content_line, style_name='바탕글')
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = '한양신명조'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '한양신명조')
        
        section_num += 1
        doc.add_paragraph()  # 섹션 간 빈 줄
    
    # 참고문헌 섹션 추가
    refs = format_references(markdown_content)
    if refs:
        doc.add_paragraph("참고 문헌", style='1.')
        for i, ref in enumerate(refs, 1):
            ref_para = doc.add_paragraph(f"[{i}] {ref}", style='바탕글')
            for run in ref_para.runs:
                run.font.size = Pt(9)
                run.font.name = '한양신명조'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '한양신명조')
    
    # 저장
    print(f"논문 저장 중: {output_path}")
    doc.save(output_path)
    print(f"완료! 논문이 {output_path}에 저장되었습니다.")
    print(f"\n적용된 완벽한 양식 요소:")
    print("  ✓ 헤더: 한양신명조 9pt")
    print("  ✓ 한국어 제목: 국문제목 스타일 (한양신명조 16pt, 굵게)")
    print("  ✓ 영문 제목: 영문제목 스타일 (한양신명조 14pt, 굵게)")
    print("  ✓ 저자 정보: 저자 스타일 (HCI Poppy)")
    print("  ✓ 요약: 요약 스타일 (한양신명조 9pt)")
    print("  ✓ 본문: 바탕글 스타일 (한양신명조 9pt)")
    print("  ✓ 섹션 제목: 1. 스타일 (한양신명조, 굵게)")
    print("  ✓ 서브섹션: 1.1 스타일 (한양견명조 9pt)")
    print("  ✓ 페이지 여백: 상하 1.2cm, 좌우 1.5cm")
    print("  ✓ 2단 컬럼 레이아웃")

if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.abspath(__file__))
    
    template_path = os.path.join(base_dir, '..', 'docs', 'ABC서식2.ABC해커톤_논문양식.docx')
    # 제출용 논문 생성
    markdown_path = os.path.join(base_dir, '..', 'docs', 'ABC해커톤_논문_제출용.md')
    output_path = os.path.join(base_dir, '..', 'docs', 'ABC해커톤_논문_제출용.docx')
    
    # 제출용 파일이 없으면 발표용 파일 사용
    if not os.path.exists(markdown_path):
        markdown_path = os.path.join(base_dir, '..', 'docs', 'ABC해커톤_논문_발표용.md')
        output_path = os.path.join(base_dir, '..', 'docs', 'ABC해커톤_논문_발표용.docx')
    
    if not os.path.exists(template_path):
        print(f"양식 파일을 찾을 수 없습니다: {template_path}")
        sys.exit(1)
    
    if not os.path.exists(markdown_path):
        print(f"마크다운 파일을 찾을 수 없습니다: {markdown_path}")
        sys.exit(1)
    
    create_perfect_paper(template_path, markdown_path, output_path)
