
import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def fix_paper(input_path, output_path):
    doc = Document(input_path)
    
    # Content to insert
    english_title = "Behavior-Based Digital Human Twin: A Machine Learning Framework for Personality Inference from Gameplay Data"
    english_author = "Jun-seok Yoo*"
    english_affiliation = "Public Human Resources Department, Hanshin University"
    
    english_abstract_title = "Abstract"
    english_abstract_text = (
        "This paper proposes a Digital Human Twin system that automatically infers and continuously learns user personality traits through machine learning-based behavior analysis. "
        "To overcome the limitations of traditional explicit survey-based personality assessments, we developed a machine learning model that predicts 4-dimensional personality weights (Logic, Intuition, Fluidity, and Complexity) by extracting implicit behavioral signals such as mouse movements and decision-making latency. "
        "The core contribution of this research is personality inference through gameplay data. By integrating behavioral data collected from various games including Minecraft, Stardew Valley, and Doki Doki Town, we construct user personality profiles and provide personalized services such as adaptive game experiences, game recommendation systems, and social matching. "
        "The system employs key AI/ML techniques: (1) Random Forest regression for personality weight prediction, (2) Exponential Moving Average (EMA) for session-to-session model adaptation, (3) time-series forecasting for future behavior patterns, (4) statistical anomaly detection for stress sensing, (5) a weighting framework for cultural bias mitigation, and (6) specialized metric processing for cross-game profile integration. "
        "Experimental results demonstrate that personality patterns can be inferred with 65-75% reliability after 5-10 sessions. The system achieved 75% parsing accuracy, 100% edge case handling, and processed over 60,000 events per second, reaching approximately 90% production readiness for real-time game data processing."
    )
    english_keywords = "Keywords: Digital Human Twin, Behavior Analysis, Game Data Collection, Online Learning, Personality Inference, Predictive Modeling, AI Ethics, Adaptive Game Experience"

    # Style mapping based on XML analysis
    TITLE_KO = "국문제목"
    TITLE_EN = "영문제목"
    AUTHOR = "저자"
    ABSTRACT = "요약"
    HEADING1 = "1."
    HEADING2 = "1.1"
    CAPTION = "그림"

    # Helper function to get style (fallback to Normal if missing)
    def get_style(doc, name):
        try:
            return doc.styles[name]
        except:
            return doc.styles['Normal']

    # 1. Global replacements and fixes
    for para in doc.paragraphs:
        # Fix double numbering
        para.text = re.sub(r"^(\d+\.)\s+(\d+\.)\s+", r"\2 ", para.text)
        para.text = re.sub(r"^(\d+\.\d+)\s+(\d+\.\d+)\s+", r"\2 ", para.text)
        
        # Fix Table Captions: "표 1. 제목" -> "<표 1> 제목"
        para.text = re.sub(r"^표\s*(\d+)\.\s*", r"<&표 \1> ", para.text).replace("<&표", "<표")
        
        # Fix Figure Captions: "그림 1. 제목" -> "(그림 1) 제목"
        para.text = re.sub(r"^그림\s*(\d+)\.\s*", r"(&그림 \1) ", para.text).replace("(&그림", "(그림")

    # 2. Insert English Sections
    found_title_idx = -1
    for i, para in enumerate(doc.paragraphs):
        if "행동 기반 디지털 휴먼 트윈" in para.text:
            found_title_idx = i
            break
            
    if found_title_idx != -1:
        # Insert English Title and Author info after Korean Title/Author
        # Typically Title is followed by Author.
        doc.paragraphs[found_title_idx].insert_paragraph_before(english_title, style=get_style(doc, TITLE_EN))
        
        # Search for Abstract to insert Before
        for i, para in enumerate(doc.paragraphs):
            if "요 약" in para.text or "초록" in para.text:
                # Insert English Abstract BEFORE the Korean one (or after, but let's follow template)
                # In the template, it was Title -> Author -> English Abstract? 
                # Actually template: Title KO -> Title EN -> Author -> Abstract KO -> Abstract EN (optional or reversed)
                # Let's just put it before 서론.
                break

    # Find Intro to insert Abstract
    for i, para in enumerate(doc.paragraphs):
        if "서론" in para.text:
            p_kw = doc.paragraphs[i].insert_paragraph_before(english_keywords, style=get_style(doc, ABSTRACT))
            p_abs = doc.paragraphs[i].insert_paragraph_before(english_abstract_text, style=get_style(doc, ABSTRACT))
            p_abs_title = doc.paragraphs[i].insert_paragraph_before(english_abstract_title, style=get_style(doc, HEADING1))
            break

    doc.save(output_path)
    print(f"Fixed paper saved to {output_path}")
    print(f"Fixed paper saved to {output_path}")

if __name__ == "__main__":
    fix_paper("docs/ABC해커톤_논문_제출용.docx", "docs/ABC해커톤_논문_제출용_fixed.docx")
