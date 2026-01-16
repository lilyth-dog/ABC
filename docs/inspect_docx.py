
from docx import Document
import sys

def check_docx(path):
    try:
        doc = Document(path)
        print(f"File: {path}")
        print(f"Paragraphs: {len(doc.paragraphs)}")
        for i, p in enumerate(doc.paragraphs[:50]): # Read first 50 paragraphs
            print(f"{i}: {p.text}")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    check_docx(sys.argv[1])
